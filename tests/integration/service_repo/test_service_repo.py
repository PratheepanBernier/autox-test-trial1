import io
import pytest

from services.ingestion import DocumentIngestionService
from services.vector_store import VectorStoreService
from models.schemas import Chunk, DocumentMetadata

class DummyEmbeddings:
    """A deterministic dummy embedding model for integration testing."""
    def embed_documents(self, texts):
        # Return a fixed-length vector for each text, based on hash for determinism
        return [[float(hash(t) % 1000)] * 5 for t in texts]
    def embed_query(self, text):
        return [float(hash(text) % 1000)] * 5

class DummyFAISS:
    """A deterministic dummy FAISS vector store for integration testing."""
    def __init__(self, documents, embeddings):
        self.docs = documents
        self.embeddings = embeddings
        self.embedded = embeddings.embed_documents([d.page_content for d in documents])
    @classmethod
    def from_documents(cls, documents, embeddings):
        return cls(documents, embeddings)
    def add_documents(self, documents):
        self.docs.extend(documents)
        self.embedded.extend(self.embeddings.embed_documents([d.page_content for d in documents]))
    def similarity_search(self, query, k=4):
        # Deterministically return the first k documents
        return self.docs[:k]
    def as_retriever(self, search_type="similarity", search_kwargs=None):
        class DummyRetriever:
            def __init__(self, docs):
                self.docs = docs
            def invoke(self, query):
                return self.docs[:search_kwargs.get("k", 4)] if search_kwargs else self.docs[:4]
        return DummyRetriever(self.docs)

@pytest.fixture
def patched_vector_store(monkeypatch):
    # Patch HuggingFaceEmbeddings and FAISS in vector_store.py
    monkeypatch.setattr("services.vector_store.HuggingFaceEmbeddings", lambda model_name: DummyEmbeddings())
    monkeypatch.setattr("services.vector_store.FAISS", DummyFAISS)
    yield

@pytest.fixture
def ingestion_service():
    return DocumentIngestionService()

@pytest.fixture
def vector_store_service(patched_vector_store):
    # Re-instantiate to use patched classes
    return VectorStoreService()

def test_ingestion_to_vector_store_txt(ingestion_service, vector_store_service):
    # Simulate a simple logistics TXT document
    txt_content = (
        "Carrier Details\n"
        "Carrier Name: Acme Logistics\n"
        "MC Number: 123456\n"
        "Driver Details\n"
        "Driver Name: John Doe\n"
        "Pickup\n"
        "123 Main St, Springfield, IL\n"
        "Drop\n"
        "456 Elm St, Shelbyville, IL\n"
        "Commodity\n"
        "Widgets\n"
        "Rate Breakdown\n"
        "Total: $1200\n"
    ).encode("utf-8")
    filename = "test_doc.txt"

    # Ingestion: process file and get chunks
    chunks = ingestion_service.process_file(txt_content, filename)
    assert isinstance(chunks, list)
    assert all(isinstance(c, Chunk) for c in chunks)
    assert any("Carrier Details" in c.text for c in chunks)
    assert any("Pickup" in c.text for c in chunks)

    # Vector store: add chunks
    vector_store_service.add_documents(chunks)
    assert vector_store_service.vector_store is not None

    # Similarity search: should return chunks with expected content
    results = vector_store_service.similarity_search("Carrier", k=2)
    assert isinstance(results, list)
    assert all(isinstance(r, Chunk) for r in results)
    assert any("Carrier" in r.text for r in results)

def test_ingestion_to_vector_store_docx(ingestion_service, vector_store_service):
    # Create a minimal DOCX file in memory
    import docx
    doc = docx.Document()
    doc.add_heading("Carrier Details", level=2)
    doc.add_paragraph("Carrier Name: Beta Transport")
    doc.add_heading("Pickup", level=2)
    doc.add_paragraph("789 Oak St, Capital City, IL")
    doc.add_heading("Drop", level=2)
    doc.add_paragraph("101 Maple Ave, Ogdenville, IL")
    fake_docx = io.BytesIO()
    doc.save(fake_docx)
    docx_bytes = fake_docx.getvalue()
    filename = "test_doc.docx"

    # Ingestion: process file and get chunks
    chunks = ingestion_service.process_file(docx_bytes, filename)
    assert isinstance(chunks, list)
    assert all(isinstance(c, Chunk) for c in chunks)
    assert any("Carrier Details" in c.text for c in chunks)
    assert any("Pickup" in c.text for c in chunks)

    # Vector store: add chunks
    vector_store_service.add_documents(chunks)
    assert vector_store_service.vector_store is not None

    # Similarity search: should return chunks with expected content
    results = vector_store_service.similarity_search("Pickup", k=2)
    assert isinstance(results, list)
    assert all(isinstance(r, Chunk) for r in results)
    assert any("Pickup" in r.text for r in results)

def test_vector_store_retriever_integration(ingestion_service, vector_store_service):
    # Add two different chunks
    chunk1 = Chunk(
        text="Carrier Details\nCarrier Name: Gamma Freight",
        metadata=DocumentMetadata(
            filename="doc1.txt", chunk_id=0, source="doc1.txt - Carrier Details", chunk_type="text"
        )
    )
    chunk2 = Chunk(
        text="Pickup\n321 Pine St, North Haverbrook, IL",
        metadata=DocumentMetadata(
            filename="doc2.txt", chunk_id=1, source="doc2.txt - Pickup", chunk_type="text"
        )
    )
    vector_store_service.add_documents([chunk1, chunk2])

    # Use retriever interface
    retriever = vector_store_service.as_retriever(search_type="similarity", search_kwargs={"k": 1})
    assert retriever is not None
    docs = retriever.invoke("Pickup location")
    assert isinstance(docs, list)
    assert len(docs) == 1
    assert "Pickup" in docs[0].page_content

def test_vector_store_empty_behavior(vector_store_service):
    # No documents added yet
    results = vector_store_service.similarity_search("anything", k=2)
    assert results == []

    retriever = vector_store_service.as_retriever()
    assert retriever is None
