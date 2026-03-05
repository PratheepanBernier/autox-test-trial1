import io
import os
import sys
import types
import pytest

import streamlit_app

from unittest.mock import patch, MagicMock, call

@pytest.fixture(autouse=True)
def patch_streamlit(monkeypatch):
    # Patch all streamlit UI calls to be no-ops or mocks
    st_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "st", st_mock)
    return st_mock

@pytest.fixture(autouse=True)
def patch_env(monkeypatch):
    # Patch environment variables for deterministic config
    monkeypatch.setenv("GROQ_API_KEY", "dummy-key")
    monkeypatch.setenv("QA_MODEL", "llama-3.3-70b-versatile")
    monkeypatch.setenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
    monkeypatch.setenv("CHUNK_SIZE", "1000")
    monkeypatch.setenv("CHUNK_OVERLAP", "200")
    monkeypatch.setenv("TOP_K", "4")

@pytest.fixture
def dummy_pdf_bytes():
    # Minimal PDF bytes for pymupdf
    return b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\ntrailer\n<<>>\n%%EOF"

@pytest.fixture
def dummy_docx_bytes():
    # Minimal docx file bytes
    from docx import Document
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Carrier Details: ACME\nRate Breakdown: $1000\nPickup: NYC\nDrop: LA\nCommodity: Widgets\nSpecial Instructions: None")
    doc.save(buf)
    return buf.getvalue()

@pytest.fixture
def dummy_txt_bytes():
    return b"Carrier Details: ACME\nRate Breakdown: $1000\nPickup: NYC\nDrop: LA\nCommodity: Widgets\nSpecial Instructions: None"

@pytest.fixture
def ingestion_service():
    return streamlit_app.DocumentIngestionService()

@pytest.fixture
def vector_store_service(monkeypatch):
    # Patch HuggingFaceEmbeddings and FAISS
    embeddings_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "HuggingFaceEmbeddings", MagicMock(return_value=embeddings_mock))
    faiss_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "FAISS", faiss_mock)
    svc = streamlit_app.VectorStoreService()
    return svc

@pytest.fixture
def rag_service(vector_store_service, monkeypatch):
    # Patch ChatGroq, ChatPromptTemplate, StrOutputParser
    llm_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "ChatGroq", MagicMock(return_value=llm_mock))
    prompt_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "ChatPromptTemplate", MagicMock())
    monkeypatch.setattr(streamlit_app.ChatPromptTemplate, "from_template", MagicMock(return_value=prompt_mock))
    str_parser_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "StrOutputParser", MagicMock(return_value=str_parser_mock))
    svc = streamlit_app.RAGService(vector_store_service)
    return svc

@pytest.fixture
def extraction_service(monkeypatch):
    llm_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "ChatGroq", MagicMock(return_value=llm_mock))
    parser_mock = MagicMock()
    parser_mock.get_format_instructions.return_value = "FORMAT"
    monkeypatch.setattr(streamlit_app, "PydanticOutputParser", MagicMock(return_value=parser_mock))
    prompt_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "ChatPromptTemplate", MagicMock())
    monkeypatch.setattr(streamlit_app.ChatPromptTemplate, "from_template", MagicMock(return_value=prompt_mock))
    svc = streamlit_app.ExtractionService()
    return svc

def test_process_file_pdf_happy_path(ingestion_service, monkeypatch, dummy_pdf_bytes):
    # Patch pymupdf.open and page.get_text
    page_mock = MagicMock()
    page_mock.get_text.return_value = "Carrier Details: ACME\nRate Breakdown: $1000"
    doc_mock = [page_mock, page_mock]
    pymupdf_mock = MagicMock()
    pymupdf_mock.open.return_value = doc_mock
    monkeypatch.setattr(streamlit_app, "pymupdf", pymupdf_mock)
    chunks = ingestion_service.process_file(dummy_pdf_bytes, "test.pdf")
    assert len(chunks) > 0
    assert all(isinstance(c, streamlit_app.Chunk) for c in chunks)
    assert any("Carrier Details" in c.text for c in chunks)
    assert all(c.metadata.filename == "test.pdf" for c in chunks)

def test_process_file_docx_happy_path(ingestion_service, dummy_docx_bytes):
    chunks = ingestion_service.process_file(dummy_docx_bytes, "test.docx")
    assert len(chunks) > 0
    assert all(isinstance(c, streamlit_app.Chunk) for c in chunks)
    assert any("Carrier Details" in c.text for c in chunks)
    assert all(c.metadata.filename == "test.docx" for c in chunks)

def test_process_file_txt_happy_path(ingestion_service, dummy_txt_bytes):
    chunks = ingestion_service.process_file(dummy_txt_bytes, "test.txt")
    assert len(chunks) > 0
    assert all(isinstance(c, streamlit_app.Chunk) for c in chunks)
    assert any("Carrier Details" in c.text for c in chunks)
    assert all(c.metadata.filename == "test.txt" for c in chunks)

def test_process_file_empty(monkeypatch, ingestion_service):
    # Should return one empty chunk or no chunks
    chunks = ingestion_service.process_file(b"", "empty.txt")
    assert isinstance(chunks, list)
    # Could be empty or a single empty chunk
    assert all(isinstance(c, streamlit_app.Chunk) for c in chunks)

def test_process_file_unsupported_extension(ingestion_service):
    # Should return empty chunks for unsupported file types
    chunks = ingestion_service.process_file(b"irrelevant", "file.xyz")
    assert isinstance(chunks, list)
    assert all(isinstance(c, streamlit_app.Chunk) for c in chunks)

def test_vector_store_add_documents_creates_store(monkeypatch):
    # Patch FAISS.from_documents
    faiss_mock = MagicMock()
    monkeypatch.setattr(streamlit_app, "FAISS", MagicMock())
    streamlit_app.FAISS.from_documents = faiss_mock
    svc = streamlit_app.VectorStoreService()
    chunk = streamlit_app.Chunk(text="foo", metadata=streamlit_app.DocumentMetadata(filename="a", chunk_id=0, source="a", chunk_type="text"))
    svc.add_documents([chunk])
    faiss_mock.assert_called_once()

def test_vector_store_add_documents_appends(monkeypatch):
    # Patch FAISS and vector_store.add_documents
    faiss_instance = MagicMock()
    faiss_cls = MagicMock()
    faiss_cls.from_documents.return_value = faiss_instance
    monkeypatch.setattr(streamlit_app, "FAISS", faiss_cls)
    svc = streamlit_app.VectorStoreService()
    chunk = streamlit_app.Chunk(text="foo", metadata=streamlit_app.DocumentMetadata(filename="a", chunk_id=0, source="a", chunk_type="text"))
    svc.add_documents([chunk])
    svc.add_documents([chunk])
    assert faiss_instance.add_documents.called

def test_rag_service_answer_happy_path(monkeypatch):
    # Patch retriever.invoke and chain.invoke
    doc_mock = MagicMock()
    doc_mock.page_content = "Carrier Details: ACME"
    doc_mock.metadata = {"filename": "a", "chunk_id": 0, "source": "a", "chunk_type": "text"}
    retriever_mock = MagicMock()
    retriever_mock.invoke.return_value = [doc_mock]
    vector_store_mock = MagicMock()
    vector_store_mock.as_retriever.return_value = retriever_mock
    vs = MagicMock()
    vs.vector_store = vector_store_mock
    llm_mock = MagicMock()
    llm_mock.invoke.return_value = "ACME is the carrier."
    prompt_mock = MagicMock()
    chain = prompt_mock.__or__.return_value.__or__.return_value
    chain.invoke.return_value = "ACME is the carrier."
    monkeypatch.setattr(streamlit_app, "ChatPromptTemplate", MagicMock())
    streamlit_app.ChatPromptTemplate.from_template.return_value = prompt_mock
    monkeypatch.setattr(streamlit_app, "ChatGroq", MagicMock(return_value=llm_mock))
    monkeypatch.setattr(streamlit_app, "StrOutputParser", MagicMock(return_value=MagicMock()))
    svc = streamlit_app.RAGService(vs)
    answer = svc.answer("Who is the carrier?")
    assert isinstance(answer, streamlit_app.SourcedAnswer)
    assert answer.confidence_score > 0.5
    assert "carrier" in answer.answer.lower()
    assert len(answer.sources) == 1

def test_rag_service_answer_not_found(monkeypatch):
    doc_mock = MagicMock()
    doc_mock.page_content = "irrelevant"
    doc_mock.metadata = {"filename": "a", "chunk_id": 0, "source": "a", "chunk_type": "text"}
    retriever_mock = MagicMock()
    retriever_mock.invoke.return_value = [doc_mock]
    vector_store_mock = MagicMock()
    vector_store_mock.as_retriever.return_value = retriever_mock
    vs = MagicMock()
    vs.vector_store = vector_store_mock
    llm_mock = MagicMock()
    llm_mock.invoke.return_value = "I cannot find the answer in the provided documents."
    prompt_mock = MagicMock()
    chain = prompt_mock.__or__.return_value.__or__.return_value
    chain.invoke.return_value = "I cannot find the answer in the provided documents."
    monkeypatch.setattr(streamlit_app, "ChatPromptTemplate", MagicMock())
    streamlit_app.ChatPromptTemplate.from_template.return_value = prompt_mock
    monkeypatch.setattr(streamlit_app, "ChatGroq", MagicMock(return_value=llm_mock))
    monkeypatch.setattr(streamlit_app, "StrOutputParser", MagicMock(return_value=MagicMock()))
    svc = streamlit_app.RAGService(vs)
    answer = svc.answer("What is the moon made of?")
    assert isinstance(answer, streamlit_app.SourcedAnswer)
    assert answer.confidence_score < 0.5
    assert "cannot find" in answer.answer.lower()

def test_extraction_service_extract_happy_path(monkeypatch):
    llm_mock = MagicMock()
    parser_mock = MagicMock()
    parser_mock.get_format_instructions.return_value = "FORMAT"
    parser_mock.invoke.return_value = streamlit_app.ShipmentData(reference_id="123", shipper="ACME", consignee="XYZ")
    prompt_mock = MagicMock()
    chain = prompt_mock.__or__.return_value.__or__.return_value
    chain.invoke.return_value = parser_mock.invoke.return_value
    monkeypatch.setattr(streamlit_app, "ChatPromptTemplate", MagicMock())
    streamlit_app.ChatPromptTemplate.from_template.return_value = prompt_mock
    monkeypatch.setattr(streamlit_app, "ChatGroq", MagicMock(return_value=llm_mock))
    monkeypatch.setattr(streamlit_app, "PydanticOutputParser", MagicMock(return_value=parser_mock))
    svc = streamlit_app.ExtractionService()
    result = svc.extract("Carrier: ACME\nConsignee: XYZ\nReference: 123")
    assert isinstance(result, streamlit_app.ShipmentData)
    assert result.reference_id == "123"
    assert result.shipper == "ACME"
    assert result.consignee == "XYZ"

def test_extraction_service_extract_missing_fields(monkeypatch):
    llm_mock = MagicMock()
    parser_mock = MagicMock()
    parser_mock.get_format_instructions.return_value = "FORMAT"
    parser_mock.invoke.return_value = streamlit_app.ShipmentData()
    prompt_mock = MagicMock()
    chain = prompt_mock.__or__.return_value.__or__.return_value
    chain.invoke.return_value = parser_mock.invoke.return_value
    monkeypatch.setattr(streamlit_app, "ChatPromptTemplate", MagicMock())
    streamlit_app.ChatPromptTemplate.from_template.return_value = prompt_mock
    monkeypatch.setattr(streamlit_app, "ChatGroq", MagicMock(return_value=llm_mock))
    monkeypatch.setattr(streamlit_app, "PydanticOutputParser", MagicMock(return_value=parser_mock))
    svc = streamlit_app.ExtractionService()
    result = svc.extract("No relevant data")
    assert isinstance(result, streamlit_app.ShipmentData)
    assert result.reference_id is None
    assert result.shipper is None
    assert result.consignee is None

def test_document_metadata_boundary_conditions():
    # Test DocumentMetadata with/without optional fields
    meta = streamlit_app.DocumentMetadata(filename="a", chunk_id=1, source="src")
    assert meta.filename == "a"
    assert meta.page_number is None
    assert meta.chunk_id == 1
    assert meta.chunk_type == "text"
    meta2 = streamlit_app.DocumentMetadata(filename="b", chunk_id=2, source="src", page_number=5, chunk_type="custom")
    assert meta2.page_number == 5
    assert meta2.chunk_type == "custom"

def test_chunk_and_sourced_answer_repr():
    meta = streamlit_app.DocumentMetadata(filename="a", chunk_id=1, source="src")
    chunk = streamlit_app.Chunk(text="foo", metadata=meta)
    answer = streamlit_app.SourcedAnswer(answer="bar", confidence_score=0.8, sources=[chunk])
    assert answer.answer == "bar"
    assert answer.confidence_score == 0.8
    assert answer.sources[0].text == "foo"

def test_shipment_data_model_fields():
    data = streamlit_app.ShipmentData(reference_id="abc", shipper="s", consignee="c")
    assert data.reference_id == "abc"
    assert data.shipper == "s"
    assert data.consignee == "c"
    # Nested models
    carrier = streamlit_app.CarrierInfo(carrier_name="CarrierX")
    data2 = streamlit_app.ShipmentData(carrier=carrier)
    assert data2.carrier.carrier_name == "CarrierX"

def test_reconciliation_equivalent_paths(monkeypatch):
    # Compare outputs of process_file for .txt and .docx with same content
    ingestion = streamlit_app.DocumentIngestionService()
    txt_bytes = b"Carrier Details: ACME\nRate Breakdown: $1000"
    # Create docx with same content
    from docx import Document
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Carrier Details: ACME\nRate Breakdown: $1000")
    doc.save(buf)
    docx_bytes = buf.getvalue()
    txt_chunks = ingestion.process_file(txt_bytes, "a.txt")
    docx_chunks = ingestion.process_file(docx_bytes, "a.docx")
    # Compare chunk texts
    txt_texts = [c.text for c in txt_chunks]
    docx_texts = [c.text for c in docx_chunks]
    assert any("Carrier Details" in t for t in txt_texts)
    assert any("Carrier Details" in t for t in docx_texts)
    # At least one chunk should be equivalent
    assert any(t in docx_texts for t in txt_texts)
