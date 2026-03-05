import pytest
from unittest.mock import patch, MagicMock
from backend.src.services.ingestion import DocumentIngestionService
from models.schemas import Chunk, DocumentMetadata

@pytest.fixture
def ingestion_service():
    return DocumentIngestionService()

@pytest.fixture
def dummy_settings(monkeypatch):
    class DummySettings:
        CHUNK_SIZE = 100
        CHUNK_OVERLAP = 10
    monkeypatch.setattr("core.config.settings", DummySettings())

@pytest.fixture
def simple_text():
    return "Carrier Details\nJohn Doe\nRate Breakdown\n$1000\nPickup\nNew York\nDrop\nLos Angeles"

@pytest.fixture
def docx_bytes():
    # Simulate a DOCX file with two paragraphs
    from io import BytesIO
    import docx
    file = BytesIO()
    d = docx.Document()
    d.add_paragraph("Carrier Details")
    d.add_paragraph("John Doe")
    d.add_paragraph("Rate Breakdown")
    d.add_paragraph("$1000")
    d.save(file)
    file.seek(0)
    return file.read()

@pytest.fixture
def pdf_bytes():
    # Simulate a PDF file with PyMuPDF
    import fitz
    from io import BytesIO
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Carrier Details\nJohn Doe\nRate Breakdown\n$1000")
    pdf_bytes = doc.write()
    doc.close()
    return pdf_bytes

def test_process_file_txt_and_docx_equivalence(ingestion_service, simple_text, docx_bytes):
    # TXT path
    txt_chunks = ingestion_service.process_file(simple_text.encode("utf-8"), "test.txt")
    # DOCX path (mock _extract_text_from_docx to return same text)
    with patch.object(DocumentIngestionService, "_extract_text_from_docx", return_value=simple_text):
        docx_chunks = ingestion_service.process_file(docx_bytes, "test.docx")
    # Reconciliation: Chunks should be equivalent in content and structure
    assert len(txt_chunks) == len(docx_chunks)
    for t, d in zip(txt_chunks, docx_chunks):
        assert t.text == d.text
        assert t.metadata.filename == d.metadata.filename
        assert t.metadata.chunk_type == d.metadata.chunk_type

def test_process_file_pdf_and_txt_equivalence(ingestion_service, simple_text, pdf_bytes):
    # PDF path (mock _extract_text_from_pdf to return same text with page marker)
    with patch.object(DocumentIngestionService, "_extract_text_from_pdf", return_value="\n\n### Page 1\n\n" + simple_text):
        pdf_chunks = ingestion_service.process_file(pdf_bytes, "test.pdf")
    # TXT path
    txt_chunks = ingestion_service.process_file(simple_text.encode("utf-8"), "test.txt")
    # Reconciliation: After cleaning, chunk text should be equivalent except for possible page marker
    # Remove page marker for comparison
    pdf_texts = [c.text.replace("### Page 1", "").strip() for c in pdf_chunks]
    txt_texts = [c.text.strip() for c in txt_chunks]
    assert pdf_texts == txt_texts

def test_process_file_unsupported_extension_returns_empty(ingestion_service):
    result = ingestion_service.process_file(b"irrelevant", "file.xyz")
    assert result == []

def test_process_file_handles_extraction_error(ingestion_service):
    # Simulate error in _extract_text_from_pdf
    with patch.object(DocumentIngestionService, "_extract_text_from_pdf", side_effect=Exception("fail")):
        result = ingestion_service.process_file(b"bad", "file.pdf")
        assert result == []

def test_add_semantic_structure_idempotence(ingestion_service, simple_text):
    # Applying twice should not double-insert headers
    once = ingestion_service._add_semantic_structure(simple_text)
    twice = ingestion_service._add_semantic_structure(once)
    assert once == twice

def test_chunk_text_with_title_grouping_consistency(ingestion_service, simple_text):
    # Should produce same output for same input regardless of grouping
    chunks1 = ingestion_service._chunk_text_with_title_grouping(simple_text, "file.txt")
    chunks2 = ingestion_service._chunk_text_with_title_grouping(simple_text, "file.txt")
    assert [c.text for c in chunks1] == [c.text for c in chunks2]
    assert [c.metadata.chunk_id for c in chunks1] == [c.metadata.chunk_id for c in chunks2]

def test_clean_chunk_removes_excess_whitespace_and_page_markers(ingestion_service):
    dirty = "\n\n\n### Page 1\n\n\nSome text\n\n\n"
    cleaned = ingestion_service._clean_chunk(dirty)
    assert "Page 1" not in cleaned
    assert cleaned == "Some text"

def test_extract_section_name_returns_general_on_no_header(ingestion_service):
    text = "No header here"
    assert ingestion_service._extract_section_name(text) == "General"

def test_extract_section_name_extracts_header(ingestion_service):
    text = "## Carrier Details\nSome content"
    assert ingestion_service._extract_section_name(text) == "Carrier Details"

def test_chunking_boundary_conditions(ingestion_service):
    # Section at the very start and end
    text = "Carrier Details\nJohn\n## Rate Breakdown\n$1000\n## Pickup\nNY"
    chunks = ingestion_service._chunk_text_with_title_grouping(text, "file.txt")
    # Should not drop any section
    assert any("Carrier Details" in c.text for c in chunks)
    assert any("Pickup" in c.text for c in chunks)

def test_chunking_empty_text_returns_empty_list(ingestion_service):
    chunks = ingestion_service._chunk_text_with_title_grouping("", "file.txt")
    assert chunks == []

def test_process_file_empty_txt_returns_empty_list(ingestion_service):
    chunks = ingestion_service.process_file(b"", "file.txt")
    assert chunks == []

def test_process_file_handles_large_input_and_overlap(ingestion_service, monkeypatch):
    # Simulate a large text and chunk splitter
    text = "## Section\n" + ("word " * 1000)
    class DummySplitter:
        def split_text(self, t):
            # Split into 2 chunks
            return [t[:100], t[100:]]
    monkeypatch.setattr(ingestion_service, "text_splitter", DummySplitter())
    chunks = ingestion_service._chunk_text_with_title_grouping(text, "file.txt")
    assert len(chunks) == 2
    assert all(isinstance(c, Chunk) for c in chunks)
    assert all(c.metadata.filename == "file.txt" for c in chunks)
    assert all(c.metadata.chunk_type == "text" for c in chunks)
    assert chunks[0].metadata.chunk_id == 0
    assert chunks[1].metadata.chunk_id == 1

def test_grouping_sections_within_same_group(ingestion_service):
    # Two carrier_info sections should be grouped
    text = "## Carrier Details\nA\n## Driver Details\nB\n## Pickup\nC"
    chunks = ingestion_service._chunk_text_with_title_grouping(text, "file.txt")
    # The first chunk should contain both Carrier and Driver Details
    assert any("Carrier Details" in c.text and "Driver Details" in c.text for c in chunks)
    # Pickup should be in a separate chunk
    assert any("Pickup" in c.text and "Carrier Details" not in c.text for c in chunks)
