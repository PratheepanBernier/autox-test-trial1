import io
import pytest
from unittest.mock import patch, MagicMock

from backend.src.services.ingestion import DocumentIngestionService
from models.schemas import Chunk, DocumentMetadata

@pytest.fixture
def service():
    return DocumentIngestionService()

@pytest.fixture
def dummy_settings(monkeypatch):
    # Patch settings for deterministic chunk size and overlap
    class DummySettings:
        CHUNK_SIZE = 10
        CHUNK_OVERLAP = 2
    monkeypatch.setattr("core.config.settings", DummySettings())

@pytest.fixture
def simple_text():
    return "Carrier Details\nJohn Doe\nRate Breakdown\n$1000\nPickup\nLocation A\nDrop\nLocation B"

@pytest.fixture
def docx_bytes():
    # Create a simple docx in memory
    import docx
    from io import BytesIO
    doc = docx.Document()
    doc.add_paragraph("Carrier Details")
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Rate Breakdown")
    doc.add_paragraph("$1000")
    doc.add_paragraph("Pickup")
    doc.add_paragraph("Location A")
    doc.add_paragraph("Drop")
    doc.add_paragraph("Location B")
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

@pytest.fixture
def pdf_bytes():
    # Mocked PDF bytes, actual content will be mocked in pymupdf
    return b"%PDF-1.4 dummy"

def make_chunk(text, filename, chunk_id, section):
    return Chunk(
        text=text,
        metadata=DocumentMetadata(
            filename=filename,
            chunk_id=chunk_id,
            source=f"{filename} - {section}",
            chunk_type="text"
        )
    )

def test_process_txt_file_happy_path(service, simple_text):
    filename = "test.txt"
    # Patch text_splitter to split by lines for deterministic output
    with patch.object(service, "text_splitter") as mock_splitter:
        mock_splitter.split_text.side_effect = lambda t: [t]
        chunks = service.process_file(simple_text.encode("utf-8"), filename)
    assert isinstance(chunks, list)
    assert all(isinstance(c, Chunk) for c in chunks)
    assert any("Carrier Details" in c.text for c in chunks)
    assert any("Rate Breakdown" in c.text for c in chunks)
    assert all(c.metadata.filename == filename for c in chunks)

def test_process_docx_file_happy_path(service, docx_bytes):
    filename = "test.docx"
    with patch.object(service, "text_splitter") as mock_splitter:
        mock_splitter.split_text.side_effect = lambda t: [t]
        chunks = service.process_file(docx_bytes, filename)
    assert isinstance(chunks, list)
    assert any("Carrier Details" in c.text for c in chunks)
    assert any("Rate Breakdown" in c.text for c in chunks)
    assert all(c.metadata.filename == filename for c in chunks)

def test_process_pdf_file_happy_path(service, pdf_bytes):
    filename = "test.pdf"
    # Patch pymupdf.open to return a mock document with two pages
    mock_page1 = MagicMock()
    mock_page1.get_text.return_value = "Carrier Details\nJohn Doe"
    mock_page2 = MagicMock()
    mock_page2.get_text.return_value = "Rate Breakdown\n$1000"
    mock_doc = [mock_page1, mock_page2]
    with patch("pymupdf.open", return_value=mock_doc):
        with patch.object(service, "text_splitter") as mock_splitter:
            mock_splitter.split_text.side_effect = lambda t: [t]
            chunks = service.process_file(pdf_bytes, filename)
    assert isinstance(chunks, list)
    assert any("Carrier Details" in c.text for c in chunks)
    assert any("Rate Breakdown" in c.text for c in chunks)
    assert all(c.metadata.filename == filename for c in chunks)

def test_process_file_unsupported_type_returns_empty(service):
    filename = "test.xlsx"
    result = service.process_file(b"dummy", filename)
    assert result == []

def test_process_file_error_handling_returns_empty(service):
    # Simulate error in _extract_text_from_pdf
    filename = "test.pdf"
    with patch.object(service, "_extract_text_from_pdf", side_effect=Exception("fail")):
        result = service.process_file(b"dummy", filename)
    assert result == []

def test_add_semantic_structure_adds_headers(service):
    text = "Carrier Details\nSome info\nRate Breakdown\n$1000"
    structured = service._add_semantic_structure(text)
    assert "## Carrier Details" in structured
    assert "## Rate Breakdown" in structured

def test_chunk_text_with_title_grouping_groups_sections(service):
    text = "\n## Carrier Details\nJohn\n## Driver Details\nJane\n## Rate Breakdown\n$1000"
    filename = "file.txt"
    # Patch text_splitter to split by section for deterministic output
    with patch.object(service, "text_splitter") as mock_splitter:
        mock_splitter.split_text.side_effect = lambda t: [t]
        chunks = service._chunk_text_with_title_grouping(text, filename)
    # Carrier Details and Driver Details should be grouped
    carrier_chunk = next((c for c in chunks if "Carrier Details" in c.text), None)
    assert carrier_chunk is not None
    assert "Driver Details" in carrier_chunk.text
    # Rate Breakdown should be its own chunk
    rate_chunk = next((c for c in chunks if "Rate Breakdown" in c.text), None)
    assert rate_chunk is not None

def test_chunk_text_with_title_grouping_handles_no_headers(service):
    text = "No headers here, just text."
    filename = "plain.txt"
    with patch.object(service, "text_splitter") as mock_splitter:
        mock_splitter.split_text.side_effect = lambda t: [t]
        chunks = service._chunk_text_with_title_grouping(text, filename)
    assert len(chunks) == 1
    assert "No headers here" in chunks[0].text

def test_clean_chunk_removes_excess_whitespace_and_page_markers(service):
    text = "\n\n\n### Page 1\n\nSome text\n\n\n### Page 2\n\n"
    cleaned = service._clean_chunk(text)
    assert "Page 1" not in cleaned
    assert "Page 2" not in cleaned
    assert cleaned.startswith("Some text")

def test_extract_section_name_with_header(service):
    text = "## Carrier Details\nSome info"
    section = service._extract_section_name(text)
    assert section == "Carrier Details"

def test_extract_section_name_without_header(service):
    text = "Just some text"
    section = service._extract_section_name(text)
    assert section == "General"

def test_reconciliation_txt_vs_docx(service, simple_text, docx_bytes):
    # Patch text_splitter to split by all text for deterministic output
    with patch.object(service, "text_splitter") as mock_splitter:
        mock_splitter.split_text.side_effect = lambda t: [t]
        txt_chunks = service.process_file(simple_text.encode("utf-8"), "file.txt")
        docx_chunks = service.process_file(docx_bytes, "file.docx")
    # The text content should be equivalent across both
    txt_texts = set(c.text.strip() for c in txt_chunks)
    docx_texts = set(c.text.strip() for c in docx_chunks)
    assert txt_texts == docx_texts

def test_reconciliation_pdf_vs_txt(service, simple_text, pdf_bytes):
    # Patch pymupdf.open to return a mock document with one page containing all text
    mock_page = MagicMock()
    mock_page.get_text.return_value = simple_text
    mock_doc = [mock_page]
    with patch("pymupdf.open", return_value=mock_doc):
        with patch.object(service, "text_splitter") as mock_splitter:
            mock_splitter.split_text.side_effect = lambda t: [t]
            pdf_chunks = service.process_file(pdf_bytes, "file.pdf")
            txt_chunks = service.process_file(simple_text.encode("utf-8"), "file.txt")
    pdf_texts = set(c.text.strip() for c in pdf_chunks)
    txt_texts = set(c.text.strip() for c in txt_chunks)
    assert pdf_texts == txt_texts

def test_reconciliation_chunking_consistency_across_paths(service, simple_text):
    # Simulate two equivalent paths: direct chunking vs. process_file
    filename = "file.txt"
    with patch.object(service, "text_splitter") as mock_splitter:
        mock_splitter.split_text.side_effect = lambda t: [t]
        # Path 1: process_file
        chunks1 = service.process_file(simple_text.encode("utf-8"), filename)
        # Path 2: manual semantic structure + chunking
        structured = service._add_semantic_structure(simple_text)
        chunks2 = service._chunk_text_with_title_grouping(structured, filename)
    texts1 = set(c.text.strip() for c in chunks1)
    texts2 = set(c.text.strip() for c in chunks2)
    assert texts1 == texts2

def test_chunk_text_with_title_grouping_boundary_conditions(service):
    # Test with empty string and minimal input
    filename = "empty.txt"
    with patch.object(service, "text_splitter") as mock_splitter:
        mock_splitter.split_text.side_effect = lambda t: [t]
        chunks = service._chunk_text_with_title_grouping("", filename)
    assert chunks == []
    with patch.object(service, "text_splitter") as mock_splitter:
        mock_splitter.split_text.side_effect = lambda t: [t]
        chunks = service._chunk_text_with_title_grouping("## OnlyHeader\n", filename)
    assert len(chunks) == 1
    assert "OnlyHeader" in chunks[0].text
