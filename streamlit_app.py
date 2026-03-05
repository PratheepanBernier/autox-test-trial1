import os
import re
import json
import logging
import sys
from typing import List, Optional, Dict, Any
from io import BytesIO

import streamlit as st
import pymupdf  # PyMuPDF
import docx
import numpy as np
from pydantic import BaseModel, Field
from dotenv import load_dotenv

# LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser, PydanticOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableParallel
from langchain_core.documents import Document

# Load environment variables
load_dotenv()

# --- Configuration (Directly from ENV or defaults) ---
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
QA_MODEL = os.getenv("QA_MODEL", "llama-3.3-70b-versatile")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 1000))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 200))
TOP_K = int(os.getenv("TOP_K", 4))

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# --- Pydantic Models (Ported from schemas.py) ---
class DocumentMetadata(BaseModel):
    filename: str
    page_number: Optional[int] = None
    chunk_id: int
    source: str
    chunk_type: str = "text"

class Chunk(BaseModel):
    text: str
    metadata: DocumentMetadata

class SourcedAnswer(BaseModel):
    answer: str
    confidence_score: float
    sources: List[Chunk]

# --- Extraction Schemas (Ported from extraction_schema.py) ---
class Location(BaseModel):
    name: Optional[str] = Field(None, description="Location name")
    address: Optional[str] = Field(None, description="Full address")
    city: Optional[str] = Field(None, description="City")
    state: Optional[str] = Field(None, description="State")
    zip_code: Optional[str] = Field(None, description="ZIP code")
    appointment_time: Optional[str] = Field(None, description="Appointment time")

class CommodityItem(BaseModel):
    commodity_name: Optional[str] = Field(None, description="Name of the commodity")
    weight: Optional[str] = Field(None, description="Weight with unit")
    quantity: Optional[str] = Field(None, description="Quantity")

class CarrierInfo(BaseModel):
    carrier_name: Optional[str] = Field(None, description="Name of the carrier company")
    mc_number: Optional[str] = Field(None, description="MC number")
    phone: Optional[str] = Field(None, description="Carrier phone number")

class DriverInfo(BaseModel):
    driver_name: Optional[str] = Field(None, description="Driver's name")
    cell_number: Optional[str] = Field(None, description="Driver's cell phone")
    truck_number: Optional[str] = Field(None, description="Truck number")

class RateInfo(BaseModel):
    total_rate: Optional[float] = Field(None, description="Total rate amount")
    currency: Optional[str] = Field(None, description="Currency code")
    rate_breakdown: Optional[Dict[str, Any]] = Field(None, description="Breakdown of rates")

class ShipmentData(BaseModel):
    reference_id: Optional[str] = Field(None, description="Reference ID or Load ID")
    shipper: Optional[str] = Field(None, description="Shipper name")
    consignee: Optional[str] = Field(None, description="Consignee name")
    carrier: Optional[CarrierInfo] = Field(None, description="Carrier information")
    driver: Optional[DriverInfo] = Field(None, description="Driver information")
    pickup: Optional[Location] = Field(None, description="Pickup details")
    drop: Optional[Location] = Field(None, description="Drop details")
    shipping_date: Optional[str] = Field(None, description="Shipping date")
    delivery_date: Optional[str] = Field(None, description="Delivery date")
    equipment_type: Optional[str] = Field(None, description="Equipment type")
    rate_info: Optional[RateInfo] = Field(None, description="Rate information")
    special_instructions: Optional[str] = Field(None, description="Special instructions")

# --- Services (Consolidated Logic) ---

class DocumentIngestionService:
    def __init__(self):
        self.section_groups = {
            'carrier_info': ['Carrier Details', 'Driver Details'],
            'customer_info': ['Customer Details', 'Shipper', 'Consignee'],
            'location_info': ['Pickup', 'Drop', 'Stops'],
            'rate_info': ['Rate Breakdown', 'Agreed Amount'],
            'commodity_info': ['Commodity', 'Description'],
            'instructions': ['Standing Instructions', 'Special Instructions']
        }
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE * 2,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n## ", "\n\n", "\n", ". ", " ", ""],
        )

    def process_file(self, file_content: bytes, filename: str) -> List[Chunk]:
        text = ""
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext == ".pdf":
            doc = pymupdf.open(stream=file_content, filetype="pdf")
            for i, page in enumerate(doc):
                text += f"\n\n### Page {i+1}\n\n" + page.get_text()
        elif file_ext == ".docx":
            doc = docx.Document(BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file_ext == ".txt":
            text = file_content.decode("utf-8")
        
        # Add semantic markers
        sections = ['Carrier Details', 'Rate Breakdown', 'Pickup', 'Drop', 'Commodity', 'Special Instructions']
        for s in sections:
            text = re.sub(f"({s})", r"\n## \1\n", text, flags=re.IGNORECASE)
        
        # Chunking
        chunks = []
        raw_chunks = self.text_splitter.split_text(text)
        for i, chunk_text in enumerate(raw_chunks):
            metadata = DocumentMetadata(
                filename=filename,
                chunk_id=i,
                source=f"{filename} - Part {i+1}",
                chunk_type="text"
            )
            chunks.append(Chunk(text=chunk_text.strip(), metadata=metadata))
        return chunks

class VectorStoreService:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
        self.vector_store = None

    def add_documents(self, chunks: List[Chunk]):
        documents = [Document(page_content=c.text, metadata=c.metadata.model_dump()) for c in chunks]
        if self.vector_store is None:
            self.vector_store = FAISS.from_documents(documents, self.embeddings)
        else:
            self.vector_store.add_documents(documents)

class RAGService:
    def __init__(self, vector_store_service):
        self.vs = vector_store_service
        self.llm = ChatGroq(temperature=0, model_name=QA_MODEL, api_key=GROQ_API_KEY)
        self.template = """
        You are a Logistics AI. Use the context to answer the question. 
        If not found, say "I cannot find the answer in the provided documents."
        CONTEXT: {context}
        QUESTION: {question}
        ANSWER:"""
        self.prompt = ChatPromptTemplate.from_template(self.template)

    def answer(self, question: str) -> SourcedAnswer:
        retriever = self.vs.vector_store.as_retriever(search_kwargs={"k": TOP_K})
        docs = retriever.invoke(question)
        
        context = "\n\n".join([d.page_content for d in docs])
        chain = self.prompt | self.llm | StrOutputParser()
        answer_text = chain.invoke({"context": context, "question": question})
        
        confidence = 0.9 if "cannot find" not in answer_text.lower() else 0.1
        sources = [Chunk(text=d.page_content, metadata=DocumentMetadata(**d.metadata)) for d in docs]
        
        return SourcedAnswer(answer=answer_text, confidence_score=confidence, sources=sources)

class ExtractionService:
    def __init__(self):
        self.llm = ChatGroq(temperature=0, model_name=QA_MODEL, api_key=GROQ_API_KEY)
        self.parser = PydanticOutputParser(pydantic_object=ShipmentData)
        self.prompt = ChatPromptTemplate.from_template(
            "Extract shipment data from text. Use null if missing.\n{format_instructions}\nText: {text}"
        )

    def extract(self, text: str) -> ShipmentData:
        chain = self.prompt | self.llm | self.parser
        return chain.invoke({"text": text, "format_instructions": self.parser.get_format_instructions()})

# --- Streamlit UI App ---

st.set_page_config(page_title="Logistics Standalone Assistant", layout="wide")

# Persistent State
if "vs" not in st.session_state:
    st.session_state.vs = VectorStoreService()
    st.session_state.ingest = DocumentIngestionService()
    st.session_state.rag = RAGService(st.session_state.vs)
    st.session_state.extractor = ExtractionService()

st.title("🚚 Logistics Document Intelligence (Standalone)")
st.info("Zero-dependency app for direct deployment. Requires GROQ_API_KEY in environment or .env file.")

if not GROQ_API_KEY:
    st.error("Missing GROQ_API_KEY in environment variables.")
    st.stop()

tabs = st.tabs(["Upload & Index", "Ask Questions", "Structured Extraction"])

with tabs[0]:
    st.header("1. Ingest Documents")
    files = st.file_uploader("Upload Logistics Docs", type=["pdf", "docx", "txt"], accept_multiple_files=True)
    if st.button("Index Documents") and files:
        with st.spinner("Processing..."):
            for f in files:
                chunks = st.session_state.ingest.process_file(f.read(), f.name)
                st.session_state.vs.add_documents(chunks)
            st.success(f"Indexed {len(files)} files successfully!")

with tabs[1]:
    st.header("2. Search & Answer")
    q = st.text_input("Ask something about your docs:")
    if st.button("Get Answer") and q:
        if st.session_state.vs.vector_store:
            with st.spinner("Analyzing..."):
                res = st.session_state.rag.answer(q)
                st.markdown(f"### Answer\n{res.answer}")
                st.metric("Confidence", f"{res.confidence_score:.2f}")
                with st.expander("Show Sources"):
                    for s in res.sources:
                        st.caption(f"Source: {s.metadata.source}")
                        st.write(s.text)
                        st.divider()
        else:
            st.warning("Please index documents first.")

with tabs[2]:
    st.header("3. Extract Data")
    f_extra = st.file_uploader("Upload for Extraction", type=["pdf", "docx", "txt"], key="extract")
    if st.button("Run Extraction") and f_extra:
        with st.spinner("Extracting..."):
            chunks = st.session_state.ingest.process_file(f_extra.read(), f_extra.name)
            full_text = "\n".join([c.text for c in chunks])
            data = st.session_state.extractor.extract(full_text)
            st.json(data.model_dump())
